"""
Este es el módulo de habilidades del asistente virtual, aquí se encuentran las funciones que se ejecutan
cuando se pide al asistente que realice alguna acción, como poner un timer, investigar algo, etc.

Es importado por el módulo principal del asistente virtual, 'src/main.py'
"""
from datetime import datetime
import threading
import requests
import os

from elevenlabs import clone, voices
from docx import Document
import spotipy

import assistant.utils as utils
from assistant.utils import TOKEN


def set_behaiviour(instruction):
    with open("src/assistant/files/persistent.txt", "w") as archivo:
        archivo.write(instruction)


def reset_behaiviour():
    with open("src/assistant/files/persistent.txt", "w") as archivo:
        archivo.write(
            "Te llamas Zaid y eres un asistente virtual, responde amigablemente")


def CrearDocumentoWord(nombreArchivo: str, tema: str, introduccion: str, desarrollo: str, conclusion: str, ruta=''):
    """
    Crea un documento de Word con el nombre especificado, el tema, la introducción, el desarrollo y la conclusión proporcionados.

    Parámetros:
    - tema (str): El tema del documento.
    - introduccion (str): La introducción del documento.
    - desarrollo (str): El desarrollo del documento.
    - conclusion (str): La conclusión del documento.
    - ruta (str, opcional): La ruta donde se guardará el archivo. Por defecto, se guarda en el directorio actual.
    """
    document = Document()  # Crear documento
    document.add_heading(tema, 0)  # Agregar título
    # Agregar contenido
    document.add_heading("Introduccion", level=1)
    document.add_paragraph(introduccion)
    document.add_heading("Desarrollo", level=1)
    document.add_paragraph(desarrollo)
    document.add_heading("Conclusión", level=1)
    document.add_paragraph(conclusion)
    # Guardar documento
    document.save(ruta + nombreArchivo+'.docx')
    # Abrir documento
    os.startfile(ruta + nombreArchivo+'.docx')


def CrearDocAux(topic: str):
    """
    Función desde la que se genera el texto relacionado con el 'topic' y se envía a la función 'CrearDocumentoWord'.

    Parámetros:
    - topic (str): El tema sobre el cual se desea crear el documento.
    """
    # Obtener links
    links = utils.getLinks(topic, 2)
    print("Doc creating links for consulting:")
    for link in links:
        print(link)
    # Obtener contexto
    context = ""
    embedd = utils.get_embedding(topic)
    print("Getting web context...")
    webContext = utils.getContext(embedd, links, 10)
    for i in webContext:
        context += i
    print("Web context getted")

    messages = [{"role": "user", "content": "information: "+context}]
    # Obtener contenido del documento
    tools = [
        {
            "type": "function",
            "function": {
                    "name": "CrearDocumentoWord",
                    "description": "Creates a word document about the given information",
                    "parameters": {
                        "type": "object",
                        "properties": {
                            "nombreArchivo": {
                                "type": "string",
                                "description": "A single-word name for the file",
                            },
                            "tema": {
                                "type": "string",
                                "description": "An adecuate title for the document",
                            },
                            "introduccion": {
                                "type": "string",
                                "description": "An adecuate introduction for the investigation",
                            },
                            "desarrollo": {
                                "type": "string",
                                "description": "An adecuate development for the investigation",
                            },
                            "conclusion": {
                                "type": "string",
                                "description": "An adecuate conclusion for the investigation",
                            }
                        },
                        "required": ["nombreArchivo", "tema", "introduccion", "desarrollo", "conclusion"],
                    }
            },
        }
    ]
    response = utils.client.chat.completions.create(
        model="gpt-3.5-turbo-1106",
        messages=messages,
        tools=tools,
        tool_choice="auto",  # auto is default, but we'll be explicit
    )
    response_message = response.choices[0].message
    tool_calls = response_message.tool_calls
    try:
        f_name = tool_calls[0].function.name
        # Ejecutar función para crear documento en un hilo
        if f_name == 'CrearDocumentoWord':
            print(f"Succesfully called {f_name}")
            CrearDocumentoWord(utils.json.loads(
                tool_calls[0].function.arguments).get('nombreArchivo'), utils.json.loads(
                tool_calls[0].function.arguments).get('tema'), utils.json.loads(
                tool_calls[0].function.arguments).get('introduccion'), utils.json.loads(
                tool_calls[0].function.arguments).get('desarrollo'), utils.json.loads(
                tool_calls[0].function.arguments).get('conclusion'))
    except Exception as e:
        print(f"Error during Doc creation {e}")


def research(query: str) -> str:
    """
    Realiza una investigación basada en una consulta dada y devuelve una respuesta generada por un modelo de lenguaje.

    Parámetros:
    - query (str): La consulta de investigación.

    Retorna:
    - output (str): La respuesta generada por el modelo de lenguaje.
    """
    # Obtener links
    links = utils.getLinks(query)
    print("Doc creating links for consulting:")
    for link in links:
        print(link)
    utils.generate_audio("Dame un momento para conseguir la información", 0)
    # Obtener contexto
    embedd = utils.get_embedding(query)
    print("Getting web context...")
    webContext = utils.getContext(embedd, links)
    context = ""

    for i in webContext:
        context += i
    print("Web context getted, creating response..")

    completion = utils.client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "Responde muy brevemente."},
            {"role": "user", "content": "Dado el siguiente contexto: " +
                context + "Responde: " + query}
        ],
        temperature=0.3,
        max_tokens=100,
        stream=True
    )
    # Obtener respuesta
    collected_chunks = []
    collected_messages = []
    output = ""
    idx = 0
    # Abrir hilos para paralelizar la generación de audio
    threads = []
    events = []
    # Generar audio
    for chunk in completion:
        collected_chunks.append(chunk)  # guardar respuesta del evento
        chunk_message = chunk.choices[0].delta.content  # extraer mensaje
        if chunk_message == '.' or chunk_message == None or chunk_message == ". \n" or chunk_message == ".\n":
            collected_messages = [
                m for m in collected_messages if m is not None]
            reply_content = ''.join([m for m in collected_messages])
            reply_content = reply_content.replace('.', '')
            if (idx != 0):
                reply_content = reply_content[1:]
            c = ". "
            if chunk_message == None:
                c = ""
            output = output + reply_content + c
            collected_messages = []
            if reply_content != "":
                print(f"Audio thread created for reply:{reply_content}")
                events.append(threading.Event())
                threads.append(threading.Thread(
                    target=utils.generate_audio, args=(reply_content, idx)))
                idx = idx+1
        collected_messages.append(chunk_message)
    # Iniciar hilos para la generación de audio
    print(f"Audio threads started")
    for thread in threads:
        thread.daemon = True
        thread.start()

    return output


def PlayInSpotify(sec=0, song="spotify:track:4UOt0ObZ4Oqi7LSU1Gzzkq"):
    """
    Reproduce una canción en Spotify.

    Parámetros:
    - sec (int, opcional): El segundo en el que se desea iniciar la canción. Por defecto, se inicia desde el principio.
    - song (str, opcional): La URI de la canción que se desea reproducir. Por defecto, se reproduce una canción x.
    """
    # Hacer la solicitud
    url = 'https://api.spotify.com/v1/me/player/play'
    headers = {
        'Authorization': f'Bearer {TOKEN}',
        'Content-Type': 'application/json'
    }
    data = {
        "uris": [
            song
        ],
        "offset": {
            "position": 0
        },
        "position_ms": sec*1000
    }
    # Enviar solicitud
    response = requests.put(url, headers=headers, json=data)
    # Verificar respuesta
    if response.status_code == 204:
        print("Solicitud exitosa.")
        print(response.text)
    else:
        print(f"Error en la solicitud: {response.status_code}")
        print(response.text)


def PreviousSong():
    """
    Esta función se utiliza para reproducir la canción anterior en Spotify.
    """
    sp = spotipy.Spotify(auth=TOKEN)
    try:
        sp.previous_track(device_id=None)
    except:
        pass


def NextSong():
    """
    Esta función se utiliza para reproducir la siguiente canción en Spotify.
    """
    sp = spotipy.Spotify(auth=TOKEN)
    try:
        sp.next_track()
    except:
        pass


def SearchASong(name):
    """
    Esta función se utiliza para buscar una canción en Spotify y reproducirla.
    """
    sp = spotipy.Spotify(auth=TOKEN)
    try:
        a = sp.search(q=name, limit=1, offset=0, type='track',
                      market=None)['tracks']['items'][0]['uri']
        PlayInSpotify(0, a)
    except:
        pass


def ResumeASong():
    """
    Esta función se utiliza para reanudar la reproducción de una canción en Spotify.
    """
    sp = spotipy.Spotify(auth=TOKEN)
    try:
        sp.start_playback(device_id=None)
    except:
        pass


def PauseASong():
    """
    Esta función se utiliza para pausar la reproducción de una canción en Spotify.
    """
    sp = spotipy.Spotify(auth=TOKEN)
    try:
        sp.pause_playback(device_id=None)
    except:
        pass


def set_timer(seconds: str) -> str:
    """
    Establece un timer de 'seconds' segundos.

    Parámetros:
    - seconds (str): El número de segundos del timer.

    Retorna:
    - (str): El mensaje de confirmación de que el timer ha sido establecido.
    """
    # Se genera un hilo para ejecutar la función mientras el asistente sigue ejecutandose
    timer_thread = threading.Thread(
        target=utils.timer_timer, args=(int(seconds), ))
    timer_thread.daemon = True
    timer_thread.start()
    print(f"{seconds} seconds timer started")
    return f"Timer de {seconds} segundos establecida"


def set_alarm(hour: str) -> str:
    """
    Establece una alarma a la hora especificada.

    Parámetros:
    - hour (str): La hora en formato HH:MM, la hora debe estar en formato de 24 horas.

    Retorna:
    - (str): El mensaje de confirmación de que la alarma ha sido establecida.
    """
    # obtener hora en un formato adecuado
    hours_a, minutes_a = map(int, hour.split(':'))
    hours_b, minutes_b, secs_b = map(
        int, datetime.now().strftime("%H:%M:%S").split(":"))
    hours = hours_a - hours_b
    minutes = minutes_a - minutes_b
    seconds = hours*3600 + minutes*60 - secs_b
    # Si la hora ya pasó, se establece la alarma para el día siguiente
    if seconds < 0:
        seconds = 24*3600 + seconds
    # Se genera un hilo para ejecutar la función mientras el asistente sigue ejecutandose
    timer_thread = threading.Thread(
        target=utils.alarm_timer, args=(seconds, ))
    timer_thread.daemon = True
    timer_thread.start()
    print(f"{seconds}s alarm timer started")

    return f"Alarma establecida a las {hour}"


def clone_voice(name):
    """
    Clona una voz y la guarda en el directorio actual.

    Parámetros:
    - name (str): El nombre de la voz a clonar.

    Retorna:
    - (str): El mensaje de confirmación de que la voz ha sido clonada.
    """
    # Se graba el audio

    utils.generate_audio(
        f"A continuación, di tu nombre y un par de cosas sobre tí", 0)
    audio_data, fs = utils.record_audio(10)
    utils.save_audio(audio_data, fs)
    # Se clona la voz
    voice = clone(
        name=name.lower(),
        description="",  # Optional
        files=["output.wav"]
    )
    with open("src/assistant/files/vkey.txt", "w") as archivo:
        archivo.write(voice.voice_id)
    with open("src/assistant/files/voces.jsonl", 'a') as archivo:
        utils.json.dump(
            {"voz": voice.name, "clave": voice.voice_id}, archivo)
    # Añade un salto de línea al final para mantener el formato JSON Lines
        archivo.write('\n')
    os.remove("output.wav")
    # Se genera un audio para confirmar que la voz ha sido clonada, y se elimina el audio grabado
    if name:
        return f"Listo, ahora puedes utilizar esta voz bajo el nombre {voice.name}"
    else:
        return "Listo, si quieres guardar una voz, clonala desde la interfaz"


def select_voice(name: str) -> str:
    """
    Selecciona una voz para utilizarla en el asistente.

    Parámetros:
    - name (str): El nombre de la voz a seleccionar.

    Retorna:
    - (str): El mensaje de confirmación de que la voz ha sido seleccionada.
    """
    voice_list = voices()
    voice_labels = [voice.name.lower() for voice in voice_list]
    # Se guarda la clave de la voz seleccionada
    name = name.lower()
    if name in voice_labels:
        selected_voice_index = voice_labels.index(name.lower())
        selected_voice_id = voice_list[selected_voice_index].voice_id
        with open("src/assistant/files/vkey.txt", "w") as archivo:
            archivo.write(selected_voice_id)
        return "Listo, esta es mi nueva voz"

    return "Lo siento, la voz seleccionada no existe, puedo clonar una voz si me lo pides, recuerda especificar un nombre para la voz!"


def Salir():
    """
    Esta función se utiliza para salir del programa.
    """
    utils.generate_audio("Hasta luego", 0)
    exit()


def brain(content: str, stm: list) -> str:

    if content[-1] == '.':
        content = content[:-1]

    print(f"User_input: {content}")
    # Detecta qué es lo que quiere el usuario
    completion = utils.client.chat.completions.create(
        model="gpt-3.5-turbo-1106",
        messages=[
            {"role": "system", "content": "You are a detector for a virtual assistant named Said that can set alarms, create documents, timers, speak, interact with music, start a process to clone voices, change your own voice, act like someone, change the acting way or the answering way, you can give information, and simple answers. Briefly identify what the user wants in the full statement."},
            {"role": "user", "content": "Brevemente, identifica lo que quiere el usuario en el siguiente enunciado completo: " + content}
        ],
        temperature=0,
        max_tokens=50,
    )
    query = completion.choices[0].message.content
    print(f"First filter detection: {query}")

    # Detecta si debe o no llamar una funcion
    messages = [
        {"role": "system", "content": "You are a virtuan assistant named said"},
        {"role": "user", "content": query}]
    tools = [
        {
            "type": "function",
            "function": {
                    "name": "set_timer",
                "description": "Creates a timer",
                "parameters": {
                            "type": "object",
                            "properties": {
                                "duration": {
                                    "type": "string",
                                    "description": "The duration of the timer in seconds"
                                }
                            },
                    "required": ["duration"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                    "name": "CrearDocAux",
                "description": "Activates when words like create, crea, genera are used",
                "parameters": {
                            "type": "object",
                            "properties": {
                                "topic": {
                                    "type": "string",
                                    "description": "The topic of the document"
                                }
                            },
                    "required": ["topic"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                    "name": "set_alarm",
                "description": "Creates an alarm using the hour",
                "parameters": {
                            "type": "object",
                            "properties": {
                                "hour": {
                                    "type": "string",
                                    "description": "The hour of the alarm in format HH:MM, the hour should be in 24h format"
                                }
                            },
                    "required": ["hour"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                    "name": "select_voice",
                "description": "Select a voice by name",
                "parameters": {
                            "type": "object",
                            "properties": {
                                "name": {
                                    "type": "string",
                                    "description": "single-word name for selecting the voice"
                                }
                            },
                    "required": ["name"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                    "name": "clone_voice",
                "description": "Starts a procces that record and clone voice",
                "parameters": {
                            "type": "object",
                            "properties": {
                                "name": {
                                    "type": "string",
                                    "description": "single-word name if provided"
                                }
                            },
                }
            }
        },
        {
            "type": "function",
            "function": {
                    "name": "SearchASong",
                "description": "used when request a song in spotify, a name should be provided",
                "parameters": {
                            "type": "object",
                            "properties": {
                                "name": {
                                    "type": "string",
                                    "description": "name of the requested song"
                                }
                            },
                    "required": ["name"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                    "name": "set_behaiviour",
                "description": "used when request to act in a certain way, or answer in certain way, or act in certain way"

            }
        },
        {
            "type": "function",
            "function": {
                    "name": "ResumeASong",
                "description": "used for resume or continue the music, instructions like play, continuar, resumir, sigue"
            }
        },
        {
            "type": "function",
            "function": {
                    "name": "PauseASong",
                "description": "used for pause or quit the music, instructions like para, quita la musica, detente"
            }
        },
        {
            "type": "function",
            "function": {
                    "name": "NextSong",
                "description": "used for skip a song or go next, instructions like siguiente, reproduce la siguiente canción"
            }
        },
        {
            "type": "function",
            "function": {
                    "name": "PreviousSong",
                "description": "used for go back a song or play previous, instructions like anterior, regresa, reproduce la canción anterior"
            }
        },
        {
            "type": "function",
            "function": {
                    "name": "Salir",
                "description": "used for quit the program, instructions like salir, termina, cierra"
            }
        },
        {
            "type": "function",
            "function": {
                    "name": "reset_behaiviour",
                "description": "used for stop acting like someone, or reset the behaviour to a virtual assistant or reset the way of acting or reset the way of answering"
            }
        },
    ]

    response = utils.client.chat.completions.create(
        model="gpt-4",
        messages=messages,
        tools=tools,
        tool_choice="auto",  # auto es el valor por defecto, pero se especifica para ser explícitos
    )

    response_message = response.choices[0].message
    tool_calls = response_message.tool_calls

    output = ''

    try:
        # Revisa si se va a llamar alguna función
        f_name = tool_calls[0].function.name
        print(f_name + " called...")
        print(f"args :- {utils.json.loads(tool_calls[0].function.arguments)}")

        # Llama a la función correspondiente
        if f_name == 'set_timer':
            output = set_timer(utils.json.loads(
                tool_calls[0].function.arguments).get('duration'))
        if f_name == 'CrearDocAux':
            output = 'En unos momentos verás el documento!'
            hilo = threading.Thread(target=CrearDocAux, args=(
                utils.json.loads(
                    tool_calls[0].function.arguments).get('topic'),))
            hilo.daemon = True
            hilo.start()

        if f_name == 'set_alarm':
            output = set_alarm(utils.json.loads(
                tool_calls[0].function.arguments).get('hour'))
        if f_name == 'select_voice':
            output = select_voice(utils.json.loads(
                tool_calls[0].function.arguments).get('name'))
        if f_name == 'research':
            output = research(utils.json.loads(
                tool_calls[0].function.arguments).get('query'))
        if f_name == 'clone_voice':
            output = clone_voice(utils.json.loads(
                tool_calls[0].function.arguments).get('name'))
        if f_name == 'SearchASong':
            SearchASong(utils.json.loads(
                tool_calls[0].function.arguments).get('name'))
            output = "Listo, la canción se está reproduciendo"
        if f_name == 'set_behaiviour':
            set_behaiviour(content)
            output = "Listo, ahora actuaré de esa manera"
        if f_name == 'reset_behaiviour':
            reset_behaiviour()
            output = "Listo, ahora actuaré como un asistente virtual"
        if f_name == 'ResumeASong':
            ResumeASong()
            output = "Listo, la canción se está reproduciendo"
        if f_name == 'PauseASong':
            PauseASong()
            output = "Listo, la canción se ha pausado"
        if f_name == 'NextSong':
            NextSong()
            output = "Listo, la canción se ha cambiado"
        if f_name == 'PreviousSong':
            PreviousSong()
            output = "Listo, la canción se ha cambiado"
        if f_name == 'Salir':
            raise SystemExit

        # Genera audio si la función dió respuesta
        if output != '':
            utils.generate_audio(output, 0)
            return output

    except Exception as e:
        # genera una respuesta normal
        print("No function called, generic response...")
        with open("src/assistant/files/persistent.txt", "r") as archivo:
            persistent_instruction = archivo.read()
        completion = utils.client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "Responde en máximo 3 oraciones"},
                {"role": "user", "content": stm[0]},
                {"role": "assistant", "content": stm[1]},
                {"role": "user", "content": stm[2]},
                {"role": "assistant", "content": stm[3]},
                {"role": "user", "content": stm[4]},
                {"role": "assistant", "content": stm[5]},
                {"role": "user", "content": stm[6]},
                {"role": "assistant", "content": stm[7]},
                {"role": "user", "content": stm[8]},
                {"role": "assistant", "content": stm[9]},
                {"role": "user", "content": persistent_instruction},
                {"role": "user", "content": content}
            ],
            temperature=0.3,
            max_tokens=100,
            stream=True
        )
        collected_chunks = []
        collected_messages = []
        output = ""
        idx = 0
        # Se generan hilos para paralelizar la generación de audio
        threads = []
        events = []
        for chunk in completion:
            collected_chunks.append(chunk)  # guardar respuesta del evento
            chunk_message = chunk.choices[0].delta.content  # extraer mensaje
            if chunk_message == '.' or chunk_message == None or chunk_message == ". \n" or chunk_message == ".\n":
                collected_messages = [
                    m for m in collected_messages if m is not None]
                reply_content = ''.join([m for m in collected_messages])
                reply_content = reply_content.replace('.', '')
                if (idx != 0):
                    reply_content = reply_content[1:]
                output = output + reply_content + ". "
                collected_messages = []
                reply_content = reply_content.strip()
                if reply_content != "":
                    print(f"Audio thread created for reply:{reply_content}")
                    events.append(threading.Event())
                    threads.append(threading.Thread(
                        target=utils.generate_audio, args=(reply_content, idx, events)))
                    idx = idx+1
            collected_messages.append(chunk_message)  # guardar el mensaje
        # Iniciar hilos para la generación de audio
        print("Audio threads started...")
        for thread in threads:
            thread.daemon = True
            thread.start()
        # Se actualiza la conversación en la memoria a corto plazo
        stm.pop(0)
        stm.pop(0)
        stm.append(content)
        stm.append(output)

    print(f"Assistant response: {output}")
    return output
